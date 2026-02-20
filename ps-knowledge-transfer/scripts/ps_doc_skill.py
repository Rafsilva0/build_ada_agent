from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Fix:   pip install python-docx")
    sys.exit(1)

NAVY     = "1F3864"
BLUE     = "4472C4"
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_RGB = RGBColor(0x1F, 0x38, 0x64)
GREY_RGB = RGBColor(0x66, 0x66, 0x66)

def _set_cell_shading(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _cell_text(cell, text: str, bold: bool = False,
               size: Pt = Pt(10), color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = color

def _make_table(doc: Document, col_widths: tuple = (2.5, 4.5)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
    hdr = table.rows[0]
    _cell_text(hdr.cells[0], "Field",    bold=True, size=Pt(10), color=WHITE)
    _cell_text(hdr.cells[1], "SC Input", bold=True, size=Pt(10), color=WHITE)
    _set_cell_shading(hdr.cells[0], NAVY)
    _set_cell_shading(hdr.cells[1], NAVY)
    return table

def _section_header(table, text: str, color_hex: str = NAVY) -> None:
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    _cell_text(row.cells[0], text, bold=True, size=Pt(11), color=WHITE)
    _set_cell_shading(row.cells[0], color_hex)

def _sub_header(table, text: str, color_hex: str = BLUE) -> None:
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    _cell_text(row.cells[0], text, bold=True, size=Pt(10), color=WHITE)
    _set_cell_shading(row.cells[0], color_hex)

def _data_row(table, field: str, value: str) -> None:
    row = table.add_row()
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(field)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.name = "Arial"
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(str(value) if value else "TBD")
    r1.font.size = Pt(9)
    r1.font.name = "Arial"

def _get(d: dict, *keys, default: str = "TBD") -> str:
    for k in keys:
        if isinstance(d, dict):
            d = d.get(k, default)
        else:
            return default
    return d or default

def _fmt_contacts(contacts) -> str:
    if not contacts:
        return "TBD"
    if isinstance(contacts, str):
        return contacts
    if isinstance(contacts, list):
        return "\n".join(f"• {c}" for c in contacts)
    return "\n".join(f"• {name} — {role}" for name, role in contacts.items())

def _fmt_list(items, bullet: str = "•") -> str:
    if not items:
        return "TBD"
    if isinstance(items, str):
        return items
    return "\n".join(f"{bullet} {item}" for item in items)

def _fmt_dict(d) -> str:
    if not d:
        return "TBD"
    if isinstance(d, str):
        return d
    return "\n".join(
        f"• {k.replace('_', ' ').title()}: {v}"
        for k, v in d.items()
    )

def generate_ps_doc(
    account_name: str,
    data: dict,
    output_dir: str | Path,
    sections: list[str] | None = None,
) -> Path:
    if sections is None:
        sections = ["general"]

    acct    = data.get("account", {})
    opp     = data.get("opportunity", {})
    demo    = data.get("demo_recap", {}) or acct.get("demo_recap", {})
    granola = data.get("granola_notes", [])
    sc_name = data.get("sc_name", "SC")

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    h = doc.add_heading("Sales to Professional Services Handoff", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(account_name)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY_RGB

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}  |  SC: {sc_name}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = GREY_RGB

    doc.add_paragraph("")
    sp = doc.add_paragraph()
    sp.add_run("Sections").bold = True
    doc.add_paragraph(
        "1. General Scoping — required for all Ada deals regardless of channel",
        style="List Number",
    )
    if "email" in sections:
        doc.add_paragraph(
            "2. Email Scoping — required when implementation includes Email",
            style="List Number",
        )
    if "voice" in sections:
        n = "3" if "email" in sections else "2"
        doc.add_paragraph(
            f"{n}. Voice Scoping — required when implementation includes Voice",
            style="List Number",
        )
    doc.add_paragraph("")

    if "general" in sections:
        h1 = doc.add_heading("GENERAL SCOPING", level=1)
        for r in h1.runs:
            r.font.color.rgb = NAVY_RGB

        t = _make_table(doc)

        overview_parts: list[str] = []
        for field in ("platform", "hq", "founded", "funding"):
            val = acct.get(field, "")
            if val:
                overview_parts.append(val if field == "platform" else f"{field.title()}: {val}")
        drivers = acct.get("business_drivers", [])
        if drivers:
            overview_parts.append("\nBusiness Drivers:")
            for d in (drivers if isinstance(drivers, list) else [drivers]):
                overview_parts.append(f"  • {d}")
        client_overview = "\n".join(overview_parts) or "TBD — needs discovery notes"

        risks = acct.get("risks", [])
        risk_list = risks if isinstance(risks, list) else ([risks] if risks else [])
        next_steps = acct.get("next_steps", [])
        arch = acct.get("key_architecture", {})

        _data_row(t, "Client Overview\n\nOverview of Account + Business case with Ada", client_overview)
        _data_row(t, "SFDC Opp", opp.get("sf_url", acct.get("salesforce_url", "TBD")))
        _data_row(t, "Solution Survey", "TBD")
        _data_row(t, "Key client stakeholders & Roles", _fmt_contacts(acct.get("contacts", {})))
        _data_row(t, "Timezone", acct.get("timezone", acct.get("hq", "TBD")))

        vols = acct.get("key_volumes", {})
        _data_row(t, "Channels currently supported", _fmt_dict(vols) if vols else acct.get("current_stack", "TBD"))
        _data_row(t, "Agent Tech Stack", acct.get("current_stack", "TBD"))
        _data_row(t, "KB Readiness\n\nFormatted and ready for AI agent ingestion or updates required", acct.get("kb_readiness", "TBD — needs assessment"))

        primary_uc   = acct.get("primary_use_case", "")
        secondary_ucs = acct.get("secondary_use_cases", [])
        scope_parts: list[str] = []
        if primary_uc:
            scope_parts.append(f"Phase 1: {primary_uc}")
        for i, uc in enumerate(secondary_ucs if isinstance(secondary_ucs, list) else [secondary_ucs]):
            scope_parts.append(f"Phase {i + 2}: {uc}")
        _data_row(t, "Project Scope\n\nWhat will Phase 1 include? What will Phase 2 include?",
                  "\n\n".join(scope_parts) or "TBD")

        _data_row(t, "Expected Launch Date?", opp.get("close_date", acct.get("close_date", "TBD")))
        _data_row(t, "Success Criteria 30 days post launch", acct.get("success_criteria", "TBD"))
        _data_row(t, "Channels\n\nWhat channels will they plan to deploy on?",
                  opp.get("product_channels", "") or acct.get("channels", "TBD"))
        _data_row(t, "Language Requirements", acct.get("languages", "English"))
        _data_row(t, "APIs / Personalization / Authentication Requirements",
                  acct.get("api_requirements", _fmt_dict(arch) if arch else "TBD"))
        _data_row(t, "Segmentation Requirements", acct.get("segmentation", "TBD"))

        _data_row(t, "Product promises made to the client / FRs?", acct.get("product_promises", "TBD — review deal notes"))
        _data_row(t, "Cluster", acct.get("cluster", "• EU"))
        _data_row(t, "Number of AI Agents", acct.get("num_agents", "1"))

        _sub_header(t, "Miscellaneous")
        _data_row(t, "Enrolled in Ada Academy", "No (pre-signature)")
        _data_row(t, "Security Requirements", acct.get("security", "Not Discussed"))
        _data_row(t, "Link + invites to Demo/Sandbox instance", acct.get("demo_url", "TBD"))
        _data_row(t, "Pilot / Opt out", "Not Discussed")
        _data_row(t, "Additional Notes / Risks",
                  _fmt_list([r for r in risk_list if "LIKELY LOST" not in str(r).upper()]) or "TBD")

        if next_steps:
            doc.add_paragraph("")
            h2 = doc.add_heading("KEY NEXT STEPS", level=2)
            for r in h2.runs:
                r.font.color.rgb = NAVY_RGB
            for step in (next_steps if isinstance(next_steps, list) else [next_steps]):
                doc.add_paragraph(str(step), style="List Number")

        if granola:
            doc.add_paragraph("")
            h3 = doc.add_heading("MEETING NOTES (from Granola)", level=2)
            for r in h3.runs:
                r.font.color.rgb = NAVY_RGB
            for note in granola[:5]:
                if isinstance(note, dict):
                    title_txt   = note.get("title", "Meeting")
                    date_txt    = note.get("date", note.get("meeting_date", ""))
                    summary_txt = note.get("summary", "") or _fmt_list(note.get("key_points", []))
                    p = doc.add_paragraph()
                    run = p.add_run(title_txt)
                    run.bold = True
                    run.font.size = Pt(10)
                    if date_txt:
                        p.add_run(f" ({date_txt})").font.size = Pt(9)
                    if summary_txt:
                        doc.add_paragraph(summary_txt)
                else:
                    doc.add_paragraph(str(note))

    if "email" in sections:
        doc.add_page_break()
        h = doc.add_heading("EMAIL SCOPING", level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY_RGB
        doc.add_paragraph("")

        t2 = _make_table(doc)
        email = data.get("email_scoping", {})

        _sub_header(t2, "Email Architecture")
        _data_row(t2, "Tech Stack\n\nIs the system your agents use to receive and respond to emails the same as your chat? Name the system.",
                  email.get("tech_stack", acct.get("current_stack", "TBD")))
        _data_row(t2, "Email landscape\n\nWhich email address(es) are your customers emailing?", email.get("email_landscape", "TBD"))
        _data_row(t2, "Webform\n\nDo you have a webform or contact form on your website?", email.get("webform", "TBD"))
        _data_row(t2, "Custom / Filter Incoming Emails\n\nDo you want to limit incoming emails to specific use cases/topics?", email.get("custom_filter", "TBD"))
        _data_row(t2, "AI Agent / Human support\n\nWhich email address will the AI Agent respond as?", email.get("ai_vs_human", "TBD"))
        _data_row(t2, "Launch plan\n\nDo you require a gradual rollout?", email.get("launch_plan", "TBD"))

        _sub_header(t2, "Email Configuration")
        _data_row(t2, "Knowledge Base\n\nAny additional sources specific to email?", email.get("kb", "Same as chat"))
        _data_row(t2, "Use cases\n\nAre there any use cases unique to email vs chat/voice?", email.get("use_cases", "TBD"))
        _data_row(t2, "Workflow Mapping\n\nAny notable differences in workflows for email vs chat?", email.get("workflow_mapping", "TBD"))
        _data_row(t2, "CC Support\n\nDo you need your AI Agent to support multiple email participants?", email.get("cc_support", "TBD"))
        _data_row(t2, "Metadata\n\nDo you currently pass metadata about your customers to Ada?", email.get("metadata", "TBD"))

        _sub_header(t2, "Email Handoffs")
        _data_row(t2, "Email / Ticketing\n\nAny differences in how AI agent hands off on email?", email.get("handoff_ticketing", "TBD"))
        _data_row(t2, "Routing\n\nSpecific use cases forwarded to a separate inbox?", email.get("routing", "TBD"))

        _sub_header(t2, "Additional Requirements")
        _data_row(t2, "Authentication\n\nDo you need the AI Agent to authenticate customers via email?", email.get("authentication", "TBD"))
        _data_row(t2, "Conversation Start\n\nDo you need a workflow at the start of each email conversation?", email.get("conversation_start", "TBD"))

    if "voice" in sections:
        doc.add_page_break()
        h = doc.add_heading("VOICE SCOPING", level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY_RGB

        t3 = _make_table(doc)
        voice = data.get("voice_scoping", {})
        risks    = acct.get("risks", [])
        risk_list = risks if isinstance(risks, list) else ([risks] if risks else [])

        _sub_header(t3, "Voice Architecture")
        _data_row(t3, "Telephony Provider", voice.get("telephony_provider", "TBD"))
        _data_row(t3, "CCaaS / Agent System\n\nWhat system do your phone agents accept calls in?", voice.get("ccaas", "TBD"))
        _data_row(t3, "SIP Integration Type", voice.get("sip_type", "TBD"))
        _data_row(t3, "Current IVR\n\nCan you share an IVR map?", voice.get("ivr", "TBD"))
        _data_row(t3, "Inbound vs Outbound", voice.get("direction", "Inbound only"))
        _data_row(t3, "Call Volume", voice.get("call_volume", "TBD"))
        _data_row(t3, "Current Agent Count", voice.get("agent_count", "TBD"))
        _data_row(t3, "Missed Call Rate", voice.get("missed_call_rate", "TBD"))

        _sub_header(t3, "Voice Use Cases")
        _data_row(t3, "Primary Voice Use Case", voice.get("primary_use_case", acct.get("primary_use_case", "TBD")))
        _data_row(t3, "Call Categorization / Triage", voice.get("triage", "TBD"))
        _data_row(t3, "Secondary Voice Use Cases", voice.get("secondary_use_cases", _fmt_list(acct.get("secondary_use_cases", []))))

        _sub_header(t3, "Voice Technical Requirements")
        _data_row(t3, "APIs Required for Voice", voice.get("apis", acct.get("api_requirements", "TBD")))
        _data_row(t3, "DTMF / Dial Pad Input", voice.get("dtmf", "TBD"))
        _data_row(t3, "SMS Capabilities", voice.get("sms", "TBD"))
        _data_row(t3, "Cross-Channel Interoperability", voice.get("cross_channel", "TBD"))

        _sub_header(t3, "Voice Handoffs")
        _data_row(t3, "Handoff to Human Agents", voice.get("handoff", "TBD"))
        _data_row(t3, "Routing Requirements", voice.get("routing", "TBD"))

        _sub_header(t3, "Voice Quality & Success Criteria")
        _data_row(t3, "Voice Quality Feedback from Demo", voice.get("demo_feedback", "TBD"))
        _data_row(t3, "Success Criteria for Voice", voice.get("success_criteria", "TBD"))
        _data_row(t3, "Voice-Specific Risks",
                  _fmt_list([r for r in risk_list if "LIKELY LOST" not in str(r).upper()]) or "TBD")
        _data_row(t3, "Timeline", voice.get("timeline", opp.get("close_date", acct.get("close_date", "TBD"))))

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    slug     = account_name.lower().replace(" ", "-")
    filename = f"PS_Knowledge_Transfer_{slug}_{datetime.now().strftime('%Y-%m-%d')}.docx"
    filepath = out / filename
    doc.save(str(filepath))
    return filepath

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a PS Knowledge Transfer .docx for an account.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--account",    required=True,  help="Account name")
    parser.add_argument("--sections",   nargs="+",      default=["general"],
                        choices=["general", "email", "voice"])
    parser.add_argument("--sc-name",    default="SC",   help="SC's full name")
    parser.add_argument("--output-dir", default="./ps-knowledge-transfer")

    group = parser.add_mutually_exclusive_group()
    group.add_argument("--data-file", help="Path to JSON file with account data")
    group.add_argument("--data-json", default="{}", help="Inline JSON string")

    args = parser.parse_args()

    if args.data_file:
        data = json.loads(Path(args.data_file).read_text(encoding="utf-8"))
    else:
        data = json.loads(args.data_json)

    data["sc_name"] = args.sc_name

    filepath = generate_ps_doc(
        account_name=args.account,
        data=data,
        output_dir=args.output_dir,
        sections=args.sections,
    )
    print(f"SUCCESS: {filepath}")

if __name__ == "__main__":
    main()
